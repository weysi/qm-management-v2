from __future__ import annotations

import base64
from io import BytesIO
from pathlib import Path
import tempfile
from zipfile import ZIP_DEFLATED, ZipFile

from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from django.utils import timezone
from rest_framework.test import APIClient
from docx import Document as WordDocument

from clients.models import Client
from documents.models import (
    Handbook,
    HandbookFile,
    Placeholder,
    PlaceholderParseCache,
    PlaceholderValue,
    WorkspaceAsset,
)
from documents.services.handbook_service import autofill_placeholders_from_client


class HandbookApiTests(TestCase):
    TRANSPARENT_PNG_BASE64 = (
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR4nGNgYAAAAAMAASsJTYQAAAAASUVORK5CYII="
    )

    def setUp(self):
        self.client = APIClient()
        self.temp_dir = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp_dir.cleanup)

        self.customer = Client.objects.create(
            name="Beispiel GmbH",
            address="Musterstrasse 1",
            zip_city="12345 Berlin",
            ceo="Max Mustermann",
            qm_manager="Erika Beispiel",
            employee_count=22,
            products="Produkt A",
            services="Service B",
            industry="Maschinenbau",
        )

    def _create_handbook(self) -> str:
        response = self.client.post(
            "/api/v1/handbooks",
            {
                "customer_id": str(self.customer.id),
                "type": "ISO9001",
            },
            format="json",
        )
        self.assertEqual(response.status_code, 201)
        return response.json()["id"]

    def _transparent_png_bytes(self) -> bytes:
        return base64.b64decode(self.TRANSPARENT_PNG_BASE64)

    def _create_zip_with_docx(self, relative_path: str, document: WordDocument) -> bytes:
        stream = BytesIO()
        document.save(stream)

        archive = BytesIO()
        with ZipFile(archive, "w", compression=ZIP_DEFLATED) as bundle:
            bundle.writestr(relative_path, stream.getvalue())
        return archive.getvalue()

    def _upload_zip_bytes(self, handbook_id: str, payload: bytes):
        upload = SimpleUploadedFile("templates.zip", payload, content_type="application/zip")
        return self.client.post(
            f"/api/v1/handbooks/{handbook_id}/upload-zip",
            {"file": upload},
            format="multipart",
        )

    def _save_signature(self, handbook_id: str, data_url: str):
        return self.client.post(
            f"/api/v1/handbooks/{handbook_id}/assets/signature",
            {"data_url": data_url, "filename": "signature.png"},
            format="json",
        )

    def _upload_workspace_asset(self, handbook_id: str, *, asset_type: str, filename: str = "asset.png"):
        upload = SimpleUploadedFile(
            filename,
            self._transparent_png_bytes(),
            content_type="image/png",
        )
        return self.client.post(
            f"/api/v1/handbooks/{handbook_id}/assets",
            {"file": upload, "asset_type": asset_type},
            format="multipart",
        )

    @override_settings(DOCUMENTS_DATA_ROOT=Path("/tmp"))
    def test_upload_zip_skips_junk_and_unsafe_paths(self):
        handbook_id = self._create_handbook()

        archive = BytesIO()
        with ZipFile(archive, "w", compression=ZIP_DEFLATED) as bundle:
            bundle.writestr("__MACOSX/._ignored", b"ignored")
            bundle.writestr("../escape.txt", b"unsafe")
            bundle.writestr(".DS_Store", b"junk")
            bundle.writestr("folder/readme.txt", b"hello")

        upload = SimpleUploadedFile(
            "templates.zip",
            archive.getvalue(),
            content_type="application/zip",
        )
        response = self.client.post(
            f"/api/v1/handbooks/{handbook_id}/upload-zip",
            {"file": upload},
            format="multipart",
        )
        self.assertEqual(response.status_code, 201)

        body = response.json()
        files = body["files"]
        self.assertEqual(len(files), 1)
        self.assertEqual(files[0]["path_in_handbook"], "folder/readme.txt")
        self.assertTrue(any(item["code"] == "SKIPPED_INVALID_PATH" for item in body["warnings"]))

    @override_settings(DOCUMENTS_DATA_ROOT=Path("/tmp"))
    def test_upload_zip_parses_docx_with_malformed_token_without_failing(self):
        handbook_id = self._create_handbook()

        doc = WordDocument()
        doc.add_paragraph("Header {{REVISION, Seite 2 von 11, gueltig ab {{VALIDITY_DATE}}")
        doc.add_paragraph("Firma {{COMPANY_NAME}}")
        stream = BytesIO()
        doc.save(stream)

        archive = BytesIO()
        with ZipFile(archive, "w", compression=ZIP_DEFLATED) as bundle:
            bundle.writestr("docs/test.docx", stream.getvalue())

        upload = SimpleUploadedFile(
            "templates.zip",
            archive.getvalue(),
            content_type="application/zip",
        )
        response = self.client.post(
            f"/api/v1/handbooks/{handbook_id}/upload-zip",
            {"file": upload},
            format="multipart",
        )
        self.assertEqual(response.status_code, 201)
        file_id = response.json()["files"][0]["id"]

        placeholders_res = self.client.get(
            f"/api/v1/handbooks/{handbook_id}/files/{file_id}/placeholders"
        )
        self.assertEqual(placeholders_res.status_code, 200)
        keys = [item["key"] for item in placeholders_res.json()["placeholders"]]
        self.assertIn("validity_date", keys)
        self.assertIn("company_name", keys)

    @override_settings(DOCUMENTS_DATA_ROOT=Path("/tmp"))
    def test_upload_zip_autofills_current_date_placeholders(self):
        handbook_id = self._create_handbook()

        doc = WordDocument()
        doc.add_paragraph("Gueltig bis {{VALIDITY_DATE,DATE}}")
        doc.add_paragraph("Heute {{DATE}}")
        upload_res = self._upload_zip_bytes(
            handbook_id,
            self._create_zip_with_docx("docs/template.docx", doc),
        )
        self.assertEqual(upload_res.status_code, 201)
        file_id = upload_res.json()["files"][0]["id"]

        placeholders_res = self.client.get(
            f"/api/v1/handbooks/{handbook_id}/files/{file_id}/placeholders"
        )
        self.assertEqual(placeholders_res.status_code, 200)
        payload = placeholders_res.json()
        placeholders_by_key = {item["key"]: item for item in payload["placeholders"]}

        expected_date = timezone.localdate().strftime("%d.%m.%Y")
        self.assertEqual(placeholders_by_key["validity_date"]["value_text"], expected_date)
        self.assertTrue(placeholders_by_key["validity_date"]["resolved"])
        self.assertEqual(placeholders_by_key["date"]["value_text"], expected_date)
        self.assertTrue(placeholders_by_key["date"]["resolved"])

    @override_settings(DOCUMENTS_DATA_ROOT=Path("/tmp"))
    def test_upload_zip_reuses_placeholder_parse_cache_by_checksum(self):
        handbook_id = self._create_handbook()
        doc = WordDocument()
        doc.add_paragraph("Firma {{client.name}}")
        payload = self._create_zip_with_docx("docs/template.docx", doc)

        first_upload = self._upload_zip_bytes(handbook_id, payload)
        self.assertEqual(first_upload.status_code, 201)
        self.assertEqual(PlaceholderParseCache.objects.count(), 1)

        second_upload = self._upload_zip_bytes(handbook_id, payload)
        self.assertEqual(second_upload.status_code, 201)
        self.assertEqual(PlaceholderParseCache.objects.count(), 1)

    @override_settings(DOCUMENTS_DATA_ROOT=Path("/tmp"))
    def test_upload_zip_autofills_from_client_text_and_assets(self):
        self.customer.logo_url = f"data:image/png;base64,{self.TRANSPARENT_PNG_BASE64}"
        self.customer.signature_url = f"data:image/png;base64,{self.TRANSPARENT_PNG_BASE64}"
        self.customer.save(update_fields=["logo_url", "signature_url", "updated_at"])

        handbook_id = self._create_handbook()
        doc = WordDocument()
        doc.add_paragraph("{{client.name}}")
        doc.add_paragraph("{{company.address}}")
        doc.add_paragraph("{{employee_count}}")
        doc.add_paragraph("{{assets.logo}}")
        doc.add_paragraph("{{assets.signature}}")
        upload_res = self._upload_zip_bytes(
            handbook_id,
            self._create_zip_with_docx("docs/template.docx", doc),
        )
        self.assertEqual(upload_res.status_code, 201)
        upload_body = upload_res.json()
        file_payload = upload_body["files"][0]
        self.assertEqual(file_payload["placeholder_total"], 5)
        self.assertEqual(file_payload["placeholder_resolved"], 5)

        file_id = file_payload["id"]
        placeholders_res = self.client.get(
            f"/api/v1/handbooks/{handbook_id}/files/{file_id}/placeholders"
        )
        self.assertEqual(placeholders_res.status_code, 200)
        payload = placeholders_res.json()
        placeholders_by_key = {item["key"]: item for item in payload["placeholders"]}

        self.assertEqual(placeholders_by_key["client.name"]["value_text"], self.customer.name)
        self.assertEqual(placeholders_by_key["company.address"]["value_text"], self.customer.address)
        self.assertEqual(
            placeholders_by_key["employee_count"]["value_text"],
            str(self.customer.employee_count),
        )
        self.assertTrue(placeholders_by_key["assets.logo"]["resolved"])
        self.assertIsNotNone(placeholders_by_key["assets.logo"]["asset_id"])
        self.assertTrue(placeholders_by_key["assets.signature"]["resolved"])
        self.assertIsNotNone(placeholders_by_key["assets.signature"]["asset_id"])

        imported_values = PlaceholderValue.objects.filter(
            handbook_id=handbook_id,
            source=PlaceholderValue.Source.IMPORTED,
        )
        self.assertTrue(imported_values.exists())

        self.assertEqual(
            WorkspaceAsset.objects.filter(
                handbook_id=handbook_id,
                asset_type=WorkspaceAsset.AssetType.LOGO,
                deleted_at__isnull=True,
            ).count(),
            1,
        )
        self.assertEqual(
            WorkspaceAsset.objects.filter(
                handbook_id=handbook_id,
                asset_type=WorkspaceAsset.AssetType.SIGNATURE,
                deleted_at__isnull=True,
            ).count(),
            1,
        )

    @override_settings(DOCUMENTS_DATA_ROOT=Path("/tmp"))
    def test_workspace_asset_upload_and_delete_sync_client_profile(self):
        handbook_id = self._create_handbook()
        doc = WordDocument()
        doc.add_paragraph("{{assets.logo}}")
        upload_res = self._upload_zip_bytes(
            handbook_id,
            self._create_zip_with_docx("docs/template.docx", doc),
        )
        self.assertEqual(upload_res.status_code, 201)
        file_id = upload_res.json()["files"][0]["id"]

        upload_asset_res = self._upload_workspace_asset(
            handbook_id,
            asset_type="logo",
            filename="logo.png",
        )
        self.assertEqual(upload_asset_res.status_code, 201)

        self.customer.refresh_from_db()
        self.assertIsNotNone(self.customer.logo_url)
        self.assertTrue(str(self.customer.logo_url).startswith("data:image/png;base64,"))

        placeholders_res = self.client.get(
            f"/api/v1/handbooks/{handbook_id}/files/{file_id}/placeholders"
        )
        self.assertEqual(placeholders_res.status_code, 200)
        logo_placeholder = placeholders_res.json()["placeholders"][0]
        self.assertTrue(logo_placeholder["resolved"])
        self.assertIsNotNone(logo_placeholder["asset_id"])

        delete_res = self.client.delete(f"/api/v1/handbooks/{handbook_id}/assets/logo")
        self.assertEqual(delete_res.status_code, 200)

        self.customer.refresh_from_db()
        self.assertIsNone(self.customer.logo_url)

        placeholders_after_delete = self.client.get(
            f"/api/v1/handbooks/{handbook_id}/files/{file_id}/placeholders"
        )
        self.assertEqual(placeholders_after_delete.status_code, 200)
        logo_placeholder_after_delete = placeholders_after_delete.json()["placeholders"][0]
        self.assertFalse(logo_placeholder_after_delete["resolved"])
        self.assertIsNone(logo_placeholder_after_delete["asset_id"])

    @override_settings(DOCUMENTS_DATA_ROOT=Path("/tmp"))
    def test_autofill_does_not_override_existing_manual_placeholder_values(self):
        handbook_id = self._create_handbook()
        doc = WordDocument()
        doc.add_paragraph("{{client.name}}")
        upload_res = self._upload_zip_bytes(
            handbook_id,
            self._create_zip_with_docx("docs/template.docx", doc),
        )
        self.assertEqual(upload_res.status_code, 201)
        file_id = upload_res.json()["files"][0]["id"]

        save_res = self.client.post(
            f"/api/v1/handbooks/{handbook_id}/placeholders/save",
            {
                "file_id": file_id,
                "values": [{"key": "client.name", "value_text": "Manual Override"}],
                "source": "MANUAL",
            },
            format="json",
        )
        self.assertEqual(save_res.status_code, 200)

        handbook = Handbook.objects.get(id=handbook_id)
        autofill_placeholders_from_client(handbook=handbook)
        value = PlaceholderValue.objects.get(handbook=handbook, key="client.name")
        self.assertEqual(value.value_text, "Manual Override")
        self.assertEqual(value.source, PlaceholderValue.Source.MANUAL)

    @override_settings(DOCUMENTS_DATA_ROOT=Path("/tmp"))
    def test_export_requires_resolved_placeholders_then_exports(self):
        handbook_id = self._create_handbook()

        doc = WordDocument()
        doc.add_paragraph("Firma {{client.name}}")
        upload_res = self._upload_zip_bytes(
            handbook_id,
            self._create_zip_with_docx("docs/template.docx", doc),
        )
        self.assertEqual(upload_res.status_code, 201)
        file_id = upload_res.json()["files"][0]["id"]

        completion_initial = self.client.get(f"/api/v1/handbooks/{handbook_id}/completion")
        self.assertEqual(completion_initial.status_code, 200)
        self.assertEqual(completion_initial.json()["required_total"], 1)
        self.assertEqual(completion_initial.json()["required_resolved"], 0)
        self.assertFalse(completion_initial.json()["is_complete_required"])

        export_fail = self.client.post(f"/api/v1/handbooks/{handbook_id}/export", {}, format="json")
        self.assertEqual(export_fail.status_code, 400)
        self.assertTrue(export_fail.json().get("errors"))

        save_res = self.client.post(
            f"/api/v1/handbooks/{handbook_id}/placeholders/save",
            {
                "file_id": file_id,
                "values": [
                    {
                        "key": "client.name",
                        "value_text": "Test Kunde GmbH",
                    }
                ],
            },
            format="json",
        )
        self.assertEqual(save_res.status_code, 200)
        self.assertIsNone(save_res.json().get("snapshot"))

        completion_after_save = self.client.get(f"/api/v1/handbooks/{handbook_id}/completion")
        self.assertEqual(completion_after_save.status_code, 200)
        self.assertEqual(completion_after_save.json()["required_total"], 1)
        self.assertEqual(completion_after_save.json()["required_resolved"], 1)
        self.assertTrue(completion_after_save.json()["is_complete_required"])

        versions_before_manual = self.client.get(f"/api/v1/handbooks/{handbook_id}/versions")
        self.assertEqual(versions_before_manual.status_code, 200)
        self.assertEqual(versions_before_manual.json()["versions"], [])

        manual_create = self.client.post(
            f"/api/v1/handbooks/{handbook_id}/versions",
            {"created_by": "test", "reason": "manual_completion"},
            format="json",
        )
        self.assertEqual(manual_create.status_code, 201)
        self.assertTrue(manual_create.json()["created"])
        self.assertFalse(manual_create.json()["snapshot"]["downloadable"])

        manual_duplicate = self.client.post(
            f"/api/v1/handbooks/{handbook_id}/versions",
            {"created_by": "test", "reason": "manual_completion"},
            format="json",
        )
        self.assertEqual(manual_duplicate.status_code, 200)
        self.assertFalse(manual_duplicate.json()["created"])

        export_ok = self.client.post(f"/api/v1/handbooks/{handbook_id}/export", {}, format="json")
        self.assertEqual(export_ok.status_code, 200)
        self.assertEqual(export_ok["Content-Type"], "application/zip")

        versions_after = self.client.get(f"/api/v1/handbooks/{handbook_id}/versions")
        self.assertEqual(versions_after.status_code, 200)
        snapshots = versions_after.json()["versions"]
        downloadable = [item for item in snapshots if item.get("downloadable")]
        self.assertTrue(downloadable)

        download_res = self.client.get(
            f"/api/v1/handbooks/{handbook_id}/versions/{downloadable[0]['version_number']}/download"
        )
        self.assertEqual(download_res.status_code, 200)
        self.assertEqual(download_res["Content-Type"], "application/zip")

    @override_settings(DOCUMENTS_DATA_ROOT=Path("/tmp"))
    def test_snapshot_delete(self):
        handbook_id = self._create_handbook()

        doc = WordDocument()
        doc.add_paragraph("{{client.name}}")
        upload_res = self._upload_zip_bytes(
            handbook_id,
            self._create_zip_with_docx("docs/template.docx", doc),
        )
        self.assertEqual(upload_res.status_code, 201)
        file_id = upload_res.json()["files"][0]["id"]

        save_res = self.client.post(
            f"/api/v1/handbooks/{handbook_id}/placeholders/save",
            {
                "file_id": file_id,
                "values": [{"key": "client.name", "value_text": "A"}],
            },
            format="json",
        )
        self.assertEqual(save_res.status_code, 200)
        self.assertIsNone(save_res.json().get("snapshot"))

        create_res = self.client.post(
            f"/api/v1/handbooks/{handbook_id}/versions",
            {"created_by": "test", "reason": "manual_completion"},
            format="json",
        )
        self.assertEqual(create_res.status_code, 201)
        self.assertTrue(create_res.json()["created"])

        versions = self.client.get(f"/api/v1/handbooks/{handbook_id}/versions").json()["versions"]
        self.assertEqual(len(versions), 1)
        version_number = versions[0]["version_number"]

        delete_res = self.client.delete(
            f"/api/v1/handbooks/{handbook_id}/versions/{version_number}"
        )
        self.assertEqual(delete_res.status_code, 200)

        versions_after = self.client.get(f"/api/v1/handbooks/{handbook_id}/versions").json()["versions"]
        self.assertEqual(versions_after, [])

    @override_settings(DOCUMENTS_DATA_ROOT=Path("/tmp"))
    def test_signature_canvas_endpoint_stores_png_and_downloads_same_bytes(self):
        handbook_id = self._create_handbook()
        png_bytes = self._transparent_png_bytes()
        data_url = f"data:image/png;base64,{self.TRANSPARENT_PNG_BASE64}"

        save_res = self._save_signature(handbook_id, data_url)
        self.assertEqual(save_res.status_code, 201)
        asset = save_res.json()["asset"]
        self.assertEqual(asset["asset_type"], "signature")

        stored_path = Path(asset["file_path"])
        self.assertTrue(stored_path.exists())
        self.assertEqual(stored_path.read_bytes(), png_bytes)

        download_res = self.client.get(f"/api/v1/handbooks/{handbook_id}/assets/signature/download")
        self.assertEqual(download_res.status_code, 200)
        self.assertEqual(download_res["Cache-Control"], "no-store, no-cache, must-revalidate, max-age=0")
        self.assertEqual(download_res["Pragma"], "no-cache")
        self.assertEqual(download_res["Expires"], "0")
        downloaded = b"".join(download_res.streaming_content)
        self.assertEqual(downloaded, png_bytes)

    @override_settings(DOCUMENTS_DATA_ROOT=Path("/tmp"))
    def test_signature_injection_uses_stored_png_bytes(self):
        handbook_id = self._create_handbook()
        doc = WordDocument()
        doc.add_paragraph("{{assets.signature}}")

        upload_res = self._upload_zip_bytes(
            handbook_id,
            self._create_zip_with_docx("docs/template.docx", doc),
        )
        self.assertEqual(upload_res.status_code, 201)

        save_signature = self._save_signature(
            handbook_id,
            f"data:image/png;base64,{self.TRANSPARENT_PNG_BASE64}",
        )
        self.assertEqual(save_signature.status_code, 201)
        signature_bytes = self._transparent_png_bytes()

        completion_res = self.client.get(f"/api/v1/handbooks/{handbook_id}/completion")
        self.assertEqual(completion_res.status_code, 200)
        self.assertEqual(completion_res.json()["required_total"], 1)
        self.assertEqual(completion_res.json()["required_resolved"], 1)
        self.assertTrue(completion_res.json()["is_complete_required"])

        export_res = self.client.post(f"/api/v1/handbooks/{handbook_id}/export", {}, format="json")
        self.assertEqual(export_res.status_code, 200)
        exported_zip_bytes = b"".join(export_res.streaming_content)

        with ZipFile(BytesIO(exported_zip_bytes), "r") as export_archive:
            rendered_docx = export_archive.read("docs/template.docx")

        with ZipFile(BytesIO(rendered_docx), "r") as doc_archive:
            media_entries = [name for name in doc_archive.namelist() if name.startswith("word/media/")]
            self.assertTrue(media_entries)
            media_payloads = [doc_archive.read(name) for name in media_entries]
            self.assertIn(signature_bytes, media_payloads)

    @override_settings(DOCUMENTS_DATA_ROOT=Path("/tmp"))
    def test_client_handbook_files_endpoint_groups_uploaded_files(self):
        handbook_id = self._create_handbook()

        doc_a = WordDocument()
        doc_a.add_paragraph("{{custom.one}}")
        doc_b = WordDocument()
        doc_b.add_paragraph("{{custom.two}}")

        archive = BytesIO()
        doc_a_stream = BytesIO()
        doc_a.save(doc_a_stream)
        doc_b_stream = BytesIO()
        doc_b.save(doc_b_stream)
        with ZipFile(archive, "w", compression=ZIP_DEFLATED) as bundle:
            bundle.writestr("docs/a.docx", doc_a_stream.getvalue())
            bundle.writestr("docs/b.docx", doc_b_stream.getvalue())

        upload_res = self._upload_zip_bytes(handbook_id, archive.getvalue())
        self.assertEqual(upload_res.status_code, 201)

        grouped_res = self.client.get(f"/api/v1/clients/{self.customer.id}/handbook-files/")
        self.assertEqual(grouped_res.status_code, 200)
        payload = grouped_res.json()

        self.assertEqual(payload["client_id"], str(self.customer.id))
        self.assertEqual(len(payload["groups"]), 1)
        group = payload["groups"][0]
        self.assertEqual(group["handbook_id"], handbook_id)
        self.assertEqual(group["file_count"], 2)
        self.assertEqual(
            [item["path_in_handbook"] for item in group["files"]],
            ["docs/a.docx", "docs/b.docx"],
        )
        self.assertTrue(all(item["deletable"] for item in group["files"]))

    @override_settings(DOCUMENTS_DATA_ROOT=Path("/tmp"))
    def test_delete_handbook_removes_workspace_and_grouped_files(self):
        handbook_id = self._create_handbook()

        doc = WordDocument()
        doc.add_paragraph("{{custom.one}}")
        upload_res = self._upload_zip_bytes(
            handbook_id,
            self._create_zip_with_docx("docs/aenderung.docx", doc),
        )
        self.assertEqual(upload_res.status_code, 201)

        grouped_before = self.client.get(
            f"/api/v1/clients/{self.customer.id}/handbook-files/"
        )
        self.assertEqual(grouped_before.status_code, 200)
        self.assertEqual(len(grouped_before.json()["groups"]), 1)

        delete_res = self.client.delete(f"/api/v1/handbooks/{handbook_id}")
        self.assertEqual(delete_res.status_code, 200)
        self.assertEqual(delete_res.json()["status"], "deleted")
        self.assertFalse(Handbook.objects.filter(id=handbook_id).exists())
        self.assertEqual(
            HandbookFile.objects.filter(handbook_id=handbook_id).count(),
            0,
        )

        grouped_after = self.client.get(
            f"/api/v1/clients/{self.customer.id}/handbook-files/"
        )
        self.assertEqual(grouped_after.status_code, 200)
        self.assertEqual(grouped_after.json()["groups"], [])

    @override_settings(DOCUMENTS_DATA_ROOT=Path("/tmp"))
    def test_delete_handbook_file_cleans_orphan_values_and_updates_completion(self):
        handbook_id = self._create_handbook()

        doc_a = WordDocument()
        doc_a.add_paragraph("{{custom.one}}")
        doc_b = WordDocument()
        doc_b.add_paragraph("{{custom.two}}")

        archive = BytesIO()
        doc_a_stream = BytesIO()
        doc_a.save(doc_a_stream)
        doc_b_stream = BytesIO()
        doc_b.save(doc_b_stream)
        with ZipFile(archive, "w", compression=ZIP_DEFLATED) as bundle:
            bundle.writestr("docs/a.docx", doc_a_stream.getvalue())
            bundle.writestr("docs/b.docx", doc_b_stream.getvalue())

        upload_res = self._upload_zip_bytes(handbook_id, archive.getvalue())
        self.assertEqual(upload_res.status_code, 201)
        files = upload_res.json()["files"]
        file_a = next(item for item in files if item["path_in_handbook"] == "docs/a.docx")
        file_b = next(item for item in files if item["path_in_handbook"] == "docs/b.docx")

        save_res = self.client.post(
            f"/api/v1/handbooks/{handbook_id}/placeholders/save",
            {
                "file_id": file_a["id"],
                "values": [{"key": "custom.one", "value_text": "Alpha"}],
                "source": "MANUAL",
            },
            format="json",
        )
        self.assertEqual(save_res.status_code, 200)
        save_res = self.client.post(
            f"/api/v1/handbooks/{handbook_id}/placeholders/save",
            {
                "file_id": file_b["id"],
                "values": [{"key": "custom.two", "value_text": "Beta"}],
                "source": "MANUAL",
            },
            format="json",
        )
        self.assertEqual(save_res.status_code, 200)

        handbook_file = HandbookFile.objects.get(id=file_a["id"])
        original_path = Path(handbook_file.original_blob_ref)
        working_path = Path(handbook_file.working_blob_ref)
        self.assertTrue(original_path.exists())

        delete_res = self.client.delete(
            f"/api/v1/handbooks/{handbook_id}/files/{file_a['id']}"
        )
        self.assertEqual(delete_res.status_code, 200)
        completion = delete_res.json()["completion"]

        self.assertFalse(HandbookFile.objects.filter(id=file_a["id"]).exists())
        self.assertFalse(Placeholder.objects.filter(handbook_file_id=file_a["id"]).exists())
        self.assertFalse(
            PlaceholderValue.objects.filter(handbook_id=handbook_id, key="custom.one").exists()
        )
        self.assertTrue(
            PlaceholderValue.objects.filter(handbook_id=handbook_id, key="custom.two").exists()
        )
        self.assertEqual([item["file_id"] for item in completion["files"]], [file_b["id"]])
        self.assertFalse(original_path.exists())
        if working_path != original_path:
            self.assertFalse(working_path.exists())

    @override_settings(DOCUMENTS_DATA_ROOT=Path("/tmp"))
    def test_delete_last_handbook_file_returns_handbook_to_draft(self):
        handbook_id = self._create_handbook()
        doc = WordDocument()
        doc.add_paragraph("{{custom.one}}")

        upload_res = self._upload_zip_bytes(
            handbook_id,
            self._create_zip_with_docx("docs/template.docx", doc),
        )
        self.assertEqual(upload_res.status_code, 201)
        file_id = upload_res.json()["files"][0]["id"]

        delete_res = self.client.delete(f"/api/v1/handbooks/{handbook_id}/files/{file_id}")
        self.assertEqual(delete_res.status_code, 200)
        completion = delete_res.json()["completion"]

        handbook = Handbook.objects.get(id=handbook_id)
        self.assertEqual(handbook.status, Handbook.Status.DRAFT)
        self.assertEqual(HandbookFile.objects.filter(handbook_id=handbook_id).count(), 0)
        self.assertEqual(completion["required_total"], 0)
        self.assertEqual(completion["required_resolved"], 0)
        self.assertEqual(completion["files"], [])
