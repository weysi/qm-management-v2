from __future__ import annotations

from io import BytesIO
from pathlib import Path

from django.db import transaction

from docx import Document as DocxDocument
from documents.models import Document, DocumentVersion, RewriteAudit

from .ai_client import AiClient, AiClientError
from .storage import LocalStorage
from .token_metrics import estimate_token_count, log_token_metrics


class RewriteValidationError(ValueError):
    pass


def _extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".docx":
        doc = DocxDocument(str(path))
        lines = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        lines.append(text)
        return "\n".join(lines).strip()

    return path.read_text(encoding="utf-8", errors="ignore")


def _target_version(document: Document, target_version: int | None) -> DocumentVersion | None:
    query = document.versions.order_by("-version_number")
    if target_version is None:
        return query.first()
    return query.filter(version_number=target_version).first()


@transaction.atomic
def rewrite_document(
    *,
    document_id: str,
    instruction: str,
    target_version: int | None = None,
) -> DocumentVersion:
    document = Document.objects.filter(id=document_id, deleted_at__isnull=True).first()
    if document is None:
        raise FileNotFoundError("Document not found")

    clean_instruction = (instruction or "").strip()
    if not clean_instruction:
        raise RewriteValidationError("instruction is required")
    if len(clean_instruction) > 4000:
        raise RewriteValidationError("instruction exceeds max length")

    version = _target_version(document, target_version)
    source_path = Path(version.file_path if version else document.original_file_path)
    content = _extract_text(source_path)
    if not content.strip():
        raise RewriteValidationError("document content is empty")

    if len(content) > 200_000:
        raise RewriteValidationError("document content too large for AI rewrite")

    audit = RewriteAudit.objects.create(
        document=document,
        source_version=version,
        instruction=clean_instruction,
        ai_model="",
        success=False,
    )

    try:
        response = AiClient().rewrite(instruction=clean_instruction, content=content)
    except AiClientError as exc:
        audit.error_message = str(exc)
        audit.save(update_fields=["error_message"])
        raise RewriteValidationError(str(exc)) from exc

    next_version = (document.versions.order_by("-version_number").first().version_number if document.versions.exists() else 0) + 1
    base = Path(document.original_file_path)
    output_path = base.parent.parent / "versions" / f"v{next_version}-{base.stem}.md"
    payload = response.content.encode("utf-8")
    LocalStorage().write_bytes(output_path, payload)
    log_token_metrics(
        path=output_path,
        estimated_token_count=estimate_token_count(response.content),
    )

    created = DocumentVersion.objects.create(
        document=document,
        version_number=next_version,
        file_path=str(output_path),
        mime_type="text/markdown",
        size_bytes=output_path.stat().st_size,
        created_by=DocumentVersion.CreatedBy.AI,
        ai_prompt=clean_instruction,
        ai_model=response.model,
        metadata={"usage": response.usage, "operation": "ai_rewrite"},
    )

    audit.ai_model = response.model
    audit.success = True
    audit.save(update_fields=["ai_model", "success"])

    return created
