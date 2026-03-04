from __future__ import annotations

from copy import deepcopy
from io import BytesIO

from docx import Document as WordDocument
from docx.shared import Inches

from .asset_placeholders import collect_asset_occurrences
from .asset_resolver import asset_missing_error
from .image_utils import (
    DEFAULT_DOCX_MAX_HEIGHT_PX,
    DEFAULT_DOCX_MAX_WIDTH_PX,
    resolve_display_size,
)
from .office_asset_types import ResolvedOfficeAsset


def _iter_container_paragraphs(container):
    for paragraph in container.paragraphs:
        yield paragraph
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_container_paragraphs(cell)


def _iter_word_paragraphs(doc: WordDocument):
    yield from _iter_container_paragraphs(doc)
    for section in doc.sections:
        yield from _iter_container_paragraphs(section.header)
        yield from _iter_container_paragraphs(section.footer)


def _run_spans(paragraph) -> list[dict[str, object]]:
    spans: list[dict[str, object]] = []
    cursor = 0
    for run in paragraph.runs:
        text = run.text or ""
        start = cursor
        end = cursor + len(text)
        spans.append(
            {
                "run": run,
                "start": start,
                "end": end,
                "text": text,
            }
        )
        cursor = end
    return spans


def _find_run_index(spans: list[dict[str, object]], pos: int) -> int | None:
    for idx, span in enumerate(spans):
        start = int(span["start"])
        end = int(span["end"])
        if end <= start:
            continue
        if start <= pos < end:
            return idx
        if start == pos:
            return idx
    return None


def _copy_run_style(source_run, target_run) -> None:
    source_rpr = source_run._r.rPr
    if source_rpr is None:
        return
    target_r = target_run._r
    if target_r.rPr is not None:
        target_r.remove(target_r.rPr)
    target_r.insert(0, deepcopy(source_rpr))


def _insert_text_run_after(paragraph, anchor_run, text: str, style_source_run):
    run = paragraph.add_run(text)
    anchor_run._r.addnext(run._r)
    _copy_run_style(style_source_run, run)
    return run


def _insert_picture_run_after(paragraph, anchor_run, payload: bytes, width_in: float, height_in: float):
    run = paragraph.add_run()
    anchor_run._r.addnext(run._r)
    run.add_picture(
        BytesIO(payload),
        width=Inches(width_in),
        height=Inches(height_in),
    )
    return run


def inject_docx_assets(
    *,
    payload: bytes,
    aliases_by_key: dict[str, list[str]],
    assets_by_key: dict[str, ResolvedOfficeAsset | None],
    fail_on_missing_asset: bool,
    error_by_key: dict[str, dict[str, object]] | None = None,
) -> tuple[bytes, list[dict[str, object]], dict[str, int]]:
    doc = WordDocument(BytesIO(payload))
    errors: list[dict[str, object]] = []
    occurrences: dict[str, int] = {}
    available_keys = set(aliases_by_key.keys())
    per_key_errors = error_by_key or {}

    for paragraph in _iter_word_paragraphs(doc):
        paragraph_text = "".join(run.text or "" for run in paragraph.runs)
        if not paragraph_text:
            continue
        matches = collect_asset_occurrences(paragraph_text, available_keys)
        if not matches:
            continue

        for match in reversed(matches):
            key = str(match.key)
            token = str(match.raw)
            occurrences[key] = occurrences.get(key, 0) + 1

            spans = _run_spans(paragraph)
            if not spans:
                continue

            start_idx = _find_run_index(spans, match.start)
            end_idx = _find_run_index(spans, max(match.start, match.end - 1))
            if start_idx is None or end_idx is None:
                continue

            start_span = spans[start_idx]
            end_span = spans[end_idx]
            start_run = start_span["run"]
            end_run = end_span["run"]
            start_text = str(start_span["text"])
            end_text = str(end_span["text"])

            local_start = max(0, match.start - int(start_span["start"]))
            local_end = max(0, match.end - int(end_span["start"]))
            prefix = start_text[:local_start]
            suffix = end_text[local_end:]

            resolved = assets_by_key.get(key)
            if resolved is None:
                if fail_on_missing_asset:
                    errors.append(per_key_errors.get(key) or asset_missing_error(key))
                else:
                    if start_idx == end_idx:
                        start_run.text = prefix + token + suffix
                    else:
                        start_run.text = prefix + token
                        end_run.text = suffix
                        for idx in range(start_idx + 1, end_idx):
                            spans[idx]["run"].text = ""
                continue

            draw_w, draw_h = resolve_display_size(
                source_width=resolved.width,
                source_height=resolved.height,
                request_width=match.width_px,
                request_height=match.height_px,
                max_width=DEFAULT_DOCX_MAX_WIDTH_PX,
                max_height=DEFAULT_DOCX_MAX_HEIGHT_PX,
            )
            width_in = draw_w / 96.0
            height_in = draw_h / 96.0

            if start_idx == end_idx:
                start_run.text = prefix
                image_anchor = _insert_picture_run_after(
                    paragraph,
                    start_run,
                    resolved.payload,
                    width_in,
                    height_in,
                )
                if suffix:
                    _insert_text_run_after(paragraph, image_anchor, suffix, start_run)
            else:
                start_run.text = prefix
                end_run.text = suffix
                for idx in range(start_idx + 1, end_idx):
                    spans[idx]["run"].text = ""
                _insert_picture_run_after(
                    paragraph,
                    start_run,
                    resolved.payload,
                    width_in,
                    height_in,
                )

    out = BytesIO()
    doc.save(out)
    return out.getvalue(), errors, occurrences
